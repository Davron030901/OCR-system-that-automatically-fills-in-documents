"""Tests for template analysis, sandboxing and rendering.

The SSTI tests are the important ones. Templates arrive from users, so the
Jinja environment is an attack surface: an unsandboxed environment turns
"upload a Word file" into arbitrary code execution on the server.
"""

import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from jinja2.exceptions import SecurityError

from packages.docgen.engine import (
    TemplateRejected,
    analyze,
    build_context,
    detect_format,
    f_amount_words,
    f_date_uz,
    f_initials,
    f_pinfl_spaced,
    f_upper_uz,
    make_environment,
    render_docx,
    sanitize_office,
)
from packages.docgen.normalize_runs import coalesce_runs, extract_text
from packages.schema.models import DocumentType, ExtractionResult, IdentityDocument
from packages.schema.translit import TURNED_COMMA

FIXTURES = Path(__file__).parent / "fixtures"
SPLIT_DOCX = FIXTURES / "split_tags.docx"


def make_docx(paragraphs: list[str]) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def sample_result() -> ExtractionResult:
    r = ExtractionResult(job_id="t", doc_type=DocumentType.ID_FRONT)
    r.person.name.surname_latin.value = "Aliyev"
    r.person.name.given_name_latin.value = "Shohruh"
    r.person.pinfl.value = "31503950012345"
    r.person.birth_date.value = "1995-03-15"
    r.documents = [IdentityDocument(doc_type="id_card")]
    r.documents[0].doc_number.value = "AA1234567"
    return r


class TestSandboxSSTI:
    """Every one of these is a real server-side template injection payload."""

    @pytest.mark.parametrize("payload", [
        "{{ ''.__class__ }}",
        "{{ ''.__class__.__mro__ }}",
        "{{ ''.__class__.__mro__[1].__subclasses__() }}",
        "{{ self.__init__.__globals__ }}",
        "{{ cycler.__init__.__globals__.os.popen('id').read() }}",
        "{{ request.application.__globals__ }}",
        "{{ config.items() }}",
        "{{ ().__class__.__bases__[0].__subclasses__() }}",
        "{{ ''.__getattribute__('__class__') }}",
    ])
    def test_injection_never_leaks_anything(self, payload):
        """Either the sandbox raises, or it renders nothing useful.

        Jinja's sandbox turns a blocked attribute into an Undefined that
        raises on any further use, so a bare `{{ ''.__class__ }}` renders as
        an empty string rather than raising. What matters is that no payload
        ever reaches a real class, module or callable.
        """
        env = make_environment()
        try:
            out = env.from_string(payload).render({})
        except Exception as exc:
            assert "uid=" not in str(exc)
            return
        for leaked in ("<class", "<module", "built-in", "uid=", "posix"):
            assert leaked not in out, f"sandbox leaked {leaked!r}: {out[:100]}"

    def test_sandbox_blocks_attribute_chain(self):
        env = make_environment()
        with pytest.raises(SecurityError):
            env.from_string("{{ ''.__class__.__mro__ }}").render({})

    def test_hostile_template_rejected_at_upload(self):
        """Defence in depth: refuse the file rather than render it emptily."""
        for payload in ["{{ ''.__class__ }}",
                        "{{ cycler.__init__.__globals__ }}",
                        "{% if config %}x{% endif %}"]:
            with pytest.raises(TemplateRejected) as exc:
                analyze(make_docx([payload]))
            assert "xavfli" in str(exc.value)

    def test_legitimate_template_still_works(self):
        env = make_environment()
        out = env.from_string("{{ name|upper_uz }}").render({"name": "aliyev"})
        assert out == "ALIYEV"


class TestFilters:
    def test_date_uz(self):
        assert f_date_uz("1995-03-15") == "15-mart, 1995-yil"

    def test_date_uz_handles_garbage(self):
        assert f_date_uz("not a date") == "not a date"

    def test_upper_uz_preserves_turned_comma(self):
        out = f_upper_uz(f"o{TURNED_COMMA}zbekiston")
        assert out == f"O{TURNED_COMMA}ZBEKISTON"
        assert TURNED_COMMA in out

    def test_pinfl_spaced(self):
        assert f_pinfl_spaced("31503950012345") == "3 1503 9500 1234 5"

    def test_pinfl_spaced_passthrough_when_wrong_length(self):
        assert f_pinfl_spaced("123") == "123"

    def test_amount_words(self):
        assert f_amount_words(1500000) == "bir million besh yuz ming"
        assert f_amount_words(0) == "nol"
        assert f_amount_words(21) == "yigirma bir"

    def test_initials(self):
        assert f_initials("Aliyev Shohruh Akmal") == "Aliyev S.A."


class TestFormatDetection:
    def test_docx_by_magic_bytes(self):
        assert detect_format(make_docx(["hello"])) == "docx"

    def test_rtf_masquerading_as_doc(self):
        """Files named .doc are very often RTF. Trust bytes, not names."""
        assert detect_format(b"{\\rtf1\\ansi test}") == "rtf"

    def test_ole2_doc(self):
        assert detect_format(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 40) == "doc"

    def test_html_masquerading(self):
        assert detect_format(b"<html><body>fake doc</body></html>") == "html"

    def test_unknown(self):
        assert detect_format(b"\x00\x01\x02\x03") == "unknown"


class TestSanitization:
    def test_clean_docx_passes(self):
        data, report = sanitize_office(make_docx(["{{ person.pinfl }}"]))
        assert data
        assert not report.findings

    def test_rtf_rejected(self):
        with pytest.raises(TemplateRejected) as e:
            sanitize_office(b"{\\rtf1\\ansi}")
        assert "format" in str(e.value).lower()

    def test_macro_document_rejected(self):
        """A .docx carrying a VBA project is a .docm in disguise."""
        base = make_docx(["hello"])
        buf = io.BytesIO()
        src = zipfile.ZipFile(io.BytesIO(base))
        with zipfile.ZipFile(buf, "w") as dst:
            for i in src.infolist():
                dst.writestr(i, src.read(i.filename))
            dst.writestr("word/vbaProject.bin", b"\x00malicious")
        with pytest.raises(TemplateRejected) as e:
            sanitize_office(buf.getvalue())
        assert "makros" in str(e.value).lower()

    def test_zip_bomb_rejected(self):
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            z.writestr("[Content_Types].xml", b"<Types/>")
            z.writestr("word/document.xml", b"A" * (60 * 1024 * 1024))
        with pytest.raises(TemplateRejected) as e:
            sanitize_office(buf.getvalue())
        assert "zip bomba" in str(e.value).lower() or "katta" in str(e.value).lower()

    def test_dde_field_stripped(self):
        base = make_docx(["hello"])
        src = zipfile.ZipFile(io.BytesIO(base))
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as dst:
            for i in src.infolist():
                data = src.read(i.filename)
                if i.filename == "word/document.xml":
                    data = data.replace(b"</w:body>",
                                        b"<w:p>DDEAUTO c:\\\\windows\\\\cmd.exe</w:p></w:body>")
                dst.writestr(i, data)
        cleaned, report = sanitize_office(buf.getvalue())
        assert b"DDEAUTO" not in zipfile.ZipFile(
            io.BytesIO(cleaned)).read("word/document.xml")
        assert any("DDE" in f for f in report.findings)

    def test_oversize_rejected(self):
        with pytest.raises(TemplateRejected):
            sanitize_office(b"PK\x03\x04" + b"\x00" * (21 * 1024 * 1024))


class TestRunCoalescing:
    def test_split_tags_become_visible(self):
        raw = SPLIT_DOCX.read_bytes()
        fixed, report = coalesce_runs(raw)
        assert report.runs_merged > 0
        # These two tags were invisible to the template engine before repair.
        assert "person.pinfl" in report.repaired_tags
        assert any("birth_date" in t for t in report.repaired_tags)

    def test_text_content_unchanged(self):
        raw = SPLIT_DOCX.read_bytes()
        fixed, _ = coalesce_runs(raw)
        assert extract_text(raw) == extract_text(fixed)

    def test_idempotent(self):
        raw = SPLIT_DOCX.read_bytes()
        once, r1 = coalesce_runs(raw)
        twice, r2 = coalesce_runs(once)
        assert r2.runs_merged == 0
        assert extract_text(once) == extract_text(twice)


class TestAnalyze:
    def test_finds_all_locations(self):
        """Table cells and headers are the ones naive implementations miss."""
        _, spec = analyze(SPLIT_DOCX.read_bytes())
        paths = {v.field_path for v in spec.variables}
        assert "documents[0].doc_number" in paths     # table cell
        assert "person.address" in paths              # header

    def test_unknown_variable_gets_a_suggestion(self):
        data = make_docx(["{{ person.middle_name }}"])
        _, spec = analyze(data)
        assert spec.unknown_variables
        path, suggestion = spec.unknown_variables[0]
        assert path == "person.middle_name"
        assert suggestion and "patronymic" in suggestion, (
            f"suggestion {suggestion!r} would send a name into the wrong field")

    def test_unclosed_tag_is_an_error_not_silence(self):
        data = make_docx(["Familiya: {{ person.name.surname_latin"])
        _, spec = analyze(data)
        assert spec.errors
        assert "Yopilmagan" in spec.errors[0]

    def test_unknown_filter_reported(self):
        data = make_docx(["{{ person.pinfl|nonexistent_filter }}"])
        _, spec = analyze(data)
        assert any("Noma'lum filtr" in e for e in spec.errors)

    def test_optional_field_detected(self):
        data = make_docx(["{{ person.address|default('') }}"])
        _, spec = analyze(data)
        assert "person.address" in spec.optional_fields


class TestRender:
    def test_renders_real_data(self):
        out = render_docx(SPLIT_DOCX.read_bytes(), sample_result())
        text = extract_text(out.content)
        assert "Aliyev" in text
        assert "31503950012345" in text
        assert "AA1234567" in text

    def test_date_filter_applied(self):
        out = render_docx(SPLIT_DOCX.read_bytes(), sample_result())
        assert "15-mart, 1995-yil" in extract_text(out.content)

    def test_missing_fields_reported_not_hidden(self):
        r = sample_result()
        r.person.address.value = None
        out = render_docx(SPLIT_DOCX.read_bytes(), r)
        assert "person.address" in out.missing_fields
        assert out.warnings

    def test_empty_field_renders_as_blank_line(self):
        r = sample_result()
        r.person.address.value = None
        out = render_docx(SPLIT_DOCX.read_bytes(), r)
        assert "____" in extract_text(out.content)

    def test_broken_template_refuses_to_render(self):
        with pytest.raises(TemplateRejected):
            render_docx(make_docx(["{{ unclosed"]), sample_result())

    def test_apostrophes_normalised_on_output(self):
        r = sample_result()
        r.person.name.surname_latin.value = "G'ofurov"     # ASCII apostrophe
        out = render_docx(SPLIT_DOCX.read_bytes(), r)
        text = extract_text(out.content)
        assert f"G{TURNED_COMMA}ofurov" in text


class TestBuildContext:
    def test_field_values_flattened(self):
        ctx, missing = build_context(sample_result())
        assert ctx["person"]["pinfl"] == "31503950012345"
        assert ctx["person"]["name"]["surname_latin"] == "Aliyev"

    def test_missing_paths_listed(self):
        _, missing = build_context(sample_result())
        assert "person.address" in missing
