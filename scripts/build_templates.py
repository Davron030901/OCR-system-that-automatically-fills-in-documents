"""Generate the bundled Word templates.

The templates are BUILT rather than committed as opaque binaries. A .docx is a
zip of XML: committed by hand it cannot be reviewed, cannot be diffed, and
nobody can tell from a pull request whether a placeholder changed. Generating
them from this file makes the actual content reviewable and makes regenerating
all of them after a schema rename a one-line command.

Run: python scripts/build_templates.py

Every template is validated after generation — it must analyse without errors
and render both with complete data and with an empty extraction result. A
template that only works when every field is present is a template that breaks
on the first real document.
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
OUT_DIR = ROOT / "packages" / "docgen" / "templates"


# --------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------
def new_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
    return doc


def para(doc, text: str, *, align=None, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def title(doc, text: str):
    return para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                space_after=12)


def right(doc, text: str):
    return para(doc, text, align=WD_ALIGN_PARAGRAPH.RIGHT)


def field_table(doc, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    return table


# --------------------------------------------------------------------------
# templates
# --------------------------------------------------------------------------
def malumotnoma(path: Path) -> None:
    """Reference letter from a place of work or study."""
    doc = new_document()
    title(doc, "MA'LUMOTNOMA")
    para(doc, "№ {{ extra.reference_number|default('____') }}",
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "")
    para(doc,
         "Ushbu ma'lumotnoma {{ person.name.surname_latin }} "
         "{{ person.name.given_name_latin }} {{ person.name.patronymic_latin }}ga "
         "berildi.")
    para(doc, "")
    field_table(doc, [
        ("Tug'ilgan sanasi", "{{ person.birth_date|date_uz }}"),
        ("Tug'ilgan joyi", "{{ person.birth_place }}"),
        ("JSHSHIR", "{{ person.pinfl|pinfl_spaced }}"),
        ("Hujjat seriyasi va raqami", "{{ documents[0].doc_number }}"),
        ("Yashash manzili", "{{ person.address }}"),
    ])
    para(doc, "")
    para(doc,
         "Ma'lumotnoma {{ extra.purpose|default('talab qilingan joyga') }} "
         "taqdim etish uchun berildi.")
    para(doc, "")
    para(doc, "{{ today|date_uz }}")
    right(doc, "Rahbar _______________ "
               "{{ extra.director|default('') }}")
    doc.save(path)


def mehnat_shartnomasi(path: Path) -> None:
    """Employment contract. Exercises conditionals and amount_words."""
    doc = new_document()
    title(doc, "MEHNAT SHARTNOMASI")
    para(doc, "№ {{ extra.contract_number|default('____') }}          "
              "{{ today|date_uz }}", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "")
    para(doc,
         "Bir tomondan {{ extra.employer|default('«Namuna» MChJ') }} "
         "(bundan keyin — Ish beruvchi) va ikkinchi tomondan "
         "{{ person.name.surname_latin|upper_uz }} "
         "{{ person.name.given_name_latin|upper_uz }} "
         "{{ person.name.patronymic_latin|upper_uz }} "
         "(bundan keyin — Xodim) quyidagilar to'g'risida ushbu shartnomani "
         "tuzdilar.")
    para(doc, "")
    para(doc, "1. SHARTNOMA PREDMETI", bold=True)
    para(doc,
         "1.1. Xodim {{ extra.position|default('____________') }} lavozimiga "
         "{{ extra.start_date|default(today)|date_uz }} dan boshlab qabul "
         "qilinadi.")
    para(doc,
         "1.2. Oylik ish haqi {{ extra.salary|default(0) }} "
         "({{ extra.salary|default(0)|amount_words }}) so'm etib belgilanadi.")
    para(doc, "")
    para(doc, "2. XODIM TO'G'RISIDAGI MA'LUMOTLAR", bold=True)
    field_table(doc, [
        ("Tug'ilgan sanasi", "{{ person.birth_date|date_uz_short }}"),
        ("JSHSHIR", "{{ person.pinfl|pinfl_spaced }}"),
        ("Hujjat", "{{ documents[0].doc_number }}, "
                   "{{ documents[0].issuing_authority }} tomonidan "
                   "{{ documents[0].issue_date|date_uz_short }} da berilgan"),
        ("Manzil", "{{ person.address }}"),
    ])
    para(doc, "")
    # A conditional, so the wording agrees with the extracted sex field.
    para(doc,
         "Xodim {% if person.sex == 'F' %}o'z farzandini parvarish qilish "
         "ta'tiliga{% else %}mehnat ta'tiliga{% endif %} qonunchilikda "
         "belgilangan tartibda chiqish huquqiga ega.")
    para(doc, "")
    para(doc, "Ish beruvchi ______________          Xodim ______________")
    right(doc, "{{ person.name.surname_latin }} "
               "{{ person.name.given_name_latin|initials }}")
    doc.save(path)


def otm_arizasi(path: Path) -> None:
    """University admission application."""
    doc = new_document()
    right(doc, "{{ extra.university|default('____________ universiteti') }} "
               "rektoriga")
    right(doc, "{{ person.name.surname_latin }} "
               "{{ person.name.given_name_latin }} "
               "{{ person.name.patronymic_latin }}dan")
    right(doc, "JSHSHIR: {{ person.pinfl|pinfl_spaced }}")
    right(doc, "Manzil: {{ person.address }}")
    para(doc, "")
    title(doc, "ARIZA")
    para(doc,
         "Meni {{ extra.speciality|default('____________') }} ta'lim "
         "yo'nalishi bo'yicha {{ extra.study_form|default('kunduzgi') }} "
         "ta'lim shakliga qabul qilishingizni so'rayman.")
    para(doc, "")
    para(doc, "Ma'lumotlarim:", bold=True)
    field_table(doc, [
        ("F.I.Sh. (kirill)", "{{ person.name.surname_cyrillic }} "
                             "{{ person.name.given_name_cyrillic }}"),
        ("Tug'ilgan sanasi", "{{ person.birth_date|date_uz }}"),
        ("Tug'ilgan joyi", "{{ person.birth_place }}"),
        ("Fuqaroligi", "{{ person.citizenship }}"),
        ("Ma'lumoti", "{{ education.institution }}, "
                      "{{ education.graduation_year }}-yil"),
        ("Diplom raqami", "{{ education.diploma_number }}"),
    ])
    para(doc, "")
    para(doc,
         "Shaxsga doir ma'lumotlarim qayta ishlanishiga rozilik beraman.")
    para(doc, "")
    para(doc, "{{ today|date_uz }}          _____________ "
              "{{ person.name.surname_latin }} "
              "{{ person.name.given_name_latin|initials }}")
    doc.save(path)


def ishonchnoma(path: Path) -> None:
    """Power of attorney draft."""
    doc = new_document()
    title(doc, "ISHONCHNOMA")
    para(doc, "{{ extra.city|default('Toshkent shahri') }}          "
              "{{ today|date_uz }}", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "")
    para(doc,
         "Men, {{ person.name.surname_latin }} "
         "{{ person.name.given_name_latin }} "
         "{{ person.name.patronymic_latin }}, "
         "{{ person.birth_date|date_uz }} da tug'ilganman, "
         "{{ documents[0].doc_number }} raqamli hujjat egasi "
         "({{ documents[0].issuing_authority }} tomonidan "
         "{{ documents[0].issue_date|date_uz_short }} da berilgan), "
         "JSHSHIR {{ person.pinfl|pinfl_spaced }}, "
         "manzil: {{ person.address }},")
    para(doc, "")
    para(doc,
         "ushbu ishonchnoma bilan "
         "{{ extra.attorney_name|default('____________') }}ga "
         "{{ extra.powers|default('quyidagi harakatlarni amalga oshirish') }} "
         "huquqini ishonib topshiraman.")
    para(doc, "")
    para(doc,
         "Ishonchnoma {{ extra.valid_until|default('bir yil') }} muddatga "
         "beriladi. Ishonchnomani boshqa shaxsga topshirish huquqisiz.")
    para(doc, "")
    para(doc, "Imzo _______________ "
              "{{ person.name.surname_latin }} "
              "{{ person.name.given_name_latin|initials }}")
    para(doc, "")
    para(doc,
         "⚠️ NAMUNA. Ishonchnoma huquqiy kuchga ega bo'lishi uchun notarial "
         "tasdiqlanishi shart.", bold=True)
    doc.save(path)


def diplom_ilovasi(path: Path) -> None:
    """Diploma supplement reference. Exercises the table-row loop."""
    doc = new_document()
    title(doc, "DIPLOM ILOVASI MA'LUMOTNOMASI")
    para(doc, "")
    field_table(doc, [
        ("F.I.Sh.", "{{ person.name.surname_latin }} "
                    "{{ person.name.given_name_latin }} "
                    "{{ person.name.patronymic_latin }}"),
        ("Tug'ilgan sanasi", "{{ person.birth_date|date_uz }}"),
        ("Ta'lim muassasasi", "{{ education.institution }}"),
        ("Ta'lim yo'nalishi", "{{ education.speciality }} "
                              "({{ education.speciality_code }})"),
        ("Daraja", "{{ education.degree }}"),
        ("Diplom raqami", "{{ education.diploma_number }}"),
        ("Bitirgan yili", "{{ education.graduation_year }}"),
        ("O'rtacha ball", "{{ education.gpa }}"),
    ])
    para(doc, "")
    para(doc, "O'zlashtirilgan fanlar:", bold=True)

    # docxtpl's {%tr %} tags live in the row they repeat: the tag row itself is
    # removed at render time, so the loop body is exactly one table row.
    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    header = table.add_row().cells
    for i, text in enumerate(["Fan nomi", "Soat", "Kredit", "Baho"]):
        header[i].text = text
    loop_open = table.add_row().cells
    loop_open[0].text = "{%tr for s in education.subjects %}"
    body = table.add_row().cells
    body[0].text = "{{ s.name }}"
    body[1].text = "{{ s.hours }}"
    body[2].text = "{{ s.credits }}"
    body[3].text = "{{ s.grade }}"
    loop_close = table.add_row().cells
    loop_close[0].text = "{%tr endfor %}"

    para(doc, "")
    para(doc, "{{ today|date_uz }}          Rahbar _____________")
    doc.save(path)


TEMPLATES = {
    "malumotnoma.docx": malumotnoma,
    "mehnat_shartnomasi.docx": mehnat_shartnomasi,
    "otm_arizasi.docx": otm_arizasi,
    "ishonchnoma.docx": ishonchnoma,
    "diplom_ilovasi.docx": diplom_ilovasi,
}


def validate(path: Path) -> str:
    """Analyse and render each template, with full data and with none.

    Rendering against an EMPTY ExtractionResult is the important half. Real
    extractions are partial far more often than they are complete, and a
    template that only survives complete data fails on the first real upload.
    """
    from packages.docgen.engine import analyze, render_docx
    from packages.schema.models import (
        Education,
        ExtractionResult,
        FieldValue,
        IdentityDocument,
        Subject,
    )

    data = path.read_bytes()
    _fixed, spec = analyze(data, name=path.stem)
    if spec.errors:
        raise SystemExit(f"{path.name}: {spec.errors}")

    empty = ExtractionResult(job_id="empty")
    empty.documents.append(IdentityDocument())
    empty.education = Education()
    render_docx(data, empty, extra={})

    full = ExtractionResult(job_id="full")
    full.person.name.surname_latin = FieldValue(value="TOSHMATOV", confidence=1.0)
    full.person.name.given_name_latin = FieldValue(value="JASUR", confidence=1.0)
    full.person.name.patronymic_latin = FieldValue(value="BAXTIYOROVICH",
                                                   confidence=1.0)
    full.person.birth_date = FieldValue(value="1993-06-12", confidence=1.0)
    full.person.birth_place = FieldValue(value="Samarqand viloyati", confidence=1.0)
    full.person.pinfl = FieldValue(value="31206930045612", confidence=1.0)
    full.person.address = FieldValue(value="Toshkent shahri", confidence=1.0)
    full.person.sex = FieldValue(value="M", confidence=1.0)
    full.documents.append(IdentityDocument(
        doc_number=FieldValue(value="AC1928374", confidence=1.0),
        issuing_authority=FieldValue(value="IIB Chilonzor tumani", confidence=1.0),
        issue_date=FieldValue(value="2019-02-14", confidence=1.0),
    ))
    full.education = Education(
        institution=FieldValue(value="TATU", confidence=1.0),
        degree=FieldValue(value="bakalavr", confidence=1.0),
        speciality=FieldValue(value="Kompyuter injiniringi", confidence=1.0),
        speciality_code=FieldValue(value="60610100", confidence=1.0),
        diploma_number=FieldValue(value="1204/22", confidence=1.0),
        graduation_year=FieldValue(value="2022", confidence=1.0),
        gpa=FieldValue(value="4.1", confidence=1.0),
        subjects=[
            Subject(name=FieldValue(value="Oliy matematika", confidence=1.0),
                    hours=FieldValue(value="180", confidence=1.0),
                    credits=FieldValue(value="6", confidence=1.0),
                    grade=FieldValue(value="a'lo", confidence=1.0)),
            Subject(name=FieldValue(value="Fizika", confidence=1.0),
                    hours=FieldValue(value="120", confidence=1.0),
                    credits=FieldValue(value="4", confidence=1.0),
                    grade=FieldValue(value="yaxshi", confidence=1.0)),
        ],
    )
    result = render_docx(data, full, extra={"salary": 7500000})
    doc = Document(io.BytesIO(result.content))
    rendered = "\n".join(p.text for p in doc.paragraphs)
    rendered += "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    leftover = [tag for tag in ("{{", "{%") if tag in rendered]
    if leftover:
        raise SystemExit(f"{path.name}: unrendered tags remain: {leftover}")
    return f"{len(spec.variables)} variables, {len(spec.loops)} loops"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for name, builder in TEMPLATES.items():
        path = OUT_DIR / name
        builder(path)
        print(f"  {name:26} {validate(path)}")
    print(f"wrote {len(TEMPLATES)} templates to {OUT_DIR}")


if __name__ == "__main__":
    main()
